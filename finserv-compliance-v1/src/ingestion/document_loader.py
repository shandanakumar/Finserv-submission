"""
src/ingestion/document_loader.py

Multi-format regulatory document ingestion pipeline.
Handles PDF, DOCX, HTML, and plain text regulatory documents.
Assigns versioned metadata for audit-trail and supersession tracking.
"""
import os
import hashlib
import logging
from pathlib import Path
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional
import re

logger = logging.getLogger(__name__)


@dataclass
class RegulatoryDocument:
    """
    Represents a parsed regulatory document before chunking.
    
    Metadata fields are critical for:
    - Retrieval filtering (jurisdiction, regulation_family)
    - Version management (supersedes, status)
    - Audit trail (ingested_at, doc_hash)
    """
    content: str
    doc_id: str
    source: str                    # e.g. "RBI", "BIS", "ESMA"
    regulation_family: str         # e.g. "Basel_III", "MiFID_II", "FEMA"
    jurisdiction: str              # "IN", "EU", "US", "GLOBAL"
    version: str                   # ISO date of publication: "2024-03-15"
    effective_date: Optional[str] = None
    supersedes: Optional[str] = None  # doc_id of superseded document
    status: str = "active"         # active | superseded | draft
    title: str = ""
    ingested_at: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    doc_hash: str = ""             # SHA-256 of content for deduplication
    file_path: str = ""

    def __post_init__(self):
        if not self.doc_hash and self.content:
            self.doc_hash = hashlib.sha256(self.content.encode()).hexdigest()[:16]


class DocumentParser:
    """Parses regulatory documents from various formats into text."""

    def parse_pdf(self, file_path: str) -> str:
        from pdfminer.high_level import extract_text

        # ── Regular text extraction ───────────────────────────────
        try:
            text = extract_text(file_path)
            if not text or len(text.strip()) < 100:
                text = ""
        except Exception as e:
            logger.warning(f"pdfminer failed ({e}), trying PyPDF2")
            text = ""

        if not text:
            text = self._parse_pdf_pypdf2(file_path)

        cleaned_text = self._clean_text(text)

        # ── Table extraction ──────────────────────────────────────
        try:
            from src.ingestion.table_chunker import TableChunker
            tc     = TableChunker()
            chunks = tc.process_pdf(file_path)

            if chunks:
                table_lines = [c["text"] for c in chunks if c["text"].strip()]
                table_section = (
                    "\n\n7. Regulatory Targets and Figures\n\n"
                    + "\n".join(table_lines)
                )
                # Prepend so chunker processes tables first
                cleaned_text = table_section + "\n\n" + cleaned_text
                logger.info(
                    f"TableChunker: added {len(table_lines)} rows to {file_path}"
                )

        except Exception as e:
            logger.warning(f"TableChunker failed ({e}) — using text only")

        return cleaned_text

    def _extract_tables_as_text(self, file_path: str) -> str:
        """
        Extract tables using pdfplumber — preserves row/column associations.
        Handles multi-row headers common in RBI/BIS regulatory PDFs.
        """
        try:
            import pdfplumber
            table_texts = []

            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    tables = page.extract_tables()
                    for table in tables:
                        if not table or len(table) < 2:
                            continue

                        # Clean all cells
                        cleaned_table = []
                        for row in table:
                            cleaned_row = [
                                " ".join(str(cell).split()) if cell else ""
                                for cell in row
                            ]
                            if any(c for c in cleaned_row):
                                cleaned_table.append(cleaned_row)

                        if len(cleaned_table) < 2:
                            continue

                        table_text_lines = []

                        # Find the first column that looks like category labels
                        # (non-empty, non-numeric, appears in most rows)
                        label_col = 0
                        for col_idx in range(len(cleaned_table[0])):
                            col_vals = [
                                r[col_idx] for r in cleaned_table
                                if col_idx < len(r) and r[col_idx]
                            ]
                            if len(col_vals) >= 2:
                                label_col = col_idx
                                break

                        # Find value column — typically column 1
                        # (first column after label that has values)
                        value_col = label_col + 1
                        if value_col >= len(cleaned_table[0]):
                            value_col = label_col

                        # Extract label: value pairs
                        for row in cleaned_table:
                            if len(row) <= label_col:
                                continue
                            label = row[label_col].strip()
                            value = row[value_col].strip() if value_col < len(row) else ""

                            # Skip rows that are headers (both cells look like headers)
                            if not label or not value:
                                continue
                            if label.lower() in ("categories", "category", "sr no", "sr. no", "s.no", "particulars"):
                                continue

                            # Write as "Label: Value"
                            table_text_lines.append(f"{label}: {value}")

                        if table_text_lines:
                            table_texts.append(
                                f"[TABLE page {page_num}]\n" +
                                "\n".join(table_text_lines)
                            )

            return "\n\n".join(table_texts)

        except ImportError:
            logger.warning("pdfplumber not installed — tables extracted as raw text")
            return ""
        except Exception as e:
            logger.warning(f"Table extraction failed: {e}")
            return ""
    def _parse_pdf_pypdf2(self, file_path: str) -> str:
        try:
            import PyPDF2
            text_parts = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
            return self._clean_text("\n".join(text_parts))
        except Exception as e:
            logger.error(f"PyPDF2 parsing failed: {e}")
            raise
    def parse_pdf(self, file_path: str) -> str:
        from pdfminer.high_level import extract_text
        try:
            # Regular text extraction
            text = extract_text(file_path)
            if not text or len(text.strip()) < 100:
                text = self._parse_pdf_pypdf2(file_path)

            cleaned_text = self._clean_text(text)

            # Table extraction — appended at end
            table_text = self._extract_tables_as_text(file_path)
            if table_text:
                table_header = "\n\nREGULATORY TABLES — KEY TARGETS AND FIGURES\n\n"
                cleaned_text = table_header + table_text + "\n\n" + cleaned_text

            return cleaned_text

        except Exception as e:
            logger.warning(f"pdfminer failed ({e}), trying PyPDF2")
            return self._parse_pdf_pypdf2(file_path)

    def parse_docx(self, file_path: str) -> str:
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return self._clean_text("\n".join(paragraphs))
        except Exception as e:
            logger.error(f"DOCX parsing failed for {file_path}: {e}")
            raise

    def parse_text(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return self._clean_text(f.read())

    def _clean_text(self, text: str) -> str:
        """Clean extracted PDF text."""
        import re

        # Fix CID font encoding issues — common in Indian regulatory PDFs
        # (cid:415) = ti ligature, (cid:425) = tt, (cid:414) = fi, (cid:416) = tl
        cid_map = {
            "(cid:415)": "ti",
            "(cid:425)": "tt",
            "(cid:414)": "fi",
            "(cid:416)": "tl",
            "(cid:271)": "b",
            "(cid:410)": "t",
            "(cid:417)": "ffi",
            "(cid:312)": "k",
        }
        for cid, replacement in cid_map.items():
            text = text.replace(cid, replacement)

        # Remove any remaining (cid:XXX) patterns
        text = re.sub(r'\(cid:\d+\)', '', text)

        # Remove null bytes
        text = text.replace('\x00', '')

        # Normalize whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {3,}', ' ', text)
        text = text.strip()

        return text

class MetadataExtractor:
    """
    Extracts metadata from filename conventions and document content.
    
    Expected filename convention:
    {SOURCE}_{REGULATION}_{YEAR}_{SEQUENCE}_{VERSION}.{ext}
    Examples:
      RBI_FEMA_2024_015_v1.pdf
      BIS_Basel_III_2023_CRE_v2.pdf
      ESMA_MiFID2_2024_Art25_v1.pdf
    """

    # Maps source prefixes to full names and jurisdictions
    SOURCE_MAP = {
        "RBI": {"name": "Reserve Bank of India", "jurisdiction": "IN", "family": "RBI"},
        "BIS": {"name": "Bank for International Settlements", "jurisdiction": "GLOBAL", "family": "Basel_III"},
        "ESMA": {"name": "European Securities and Markets Authority", "jurisdiction": "EU", "family": "MiFID_II"},
        "EBA": {"name": "European Banking Authority", "jurisdiction": "EU", "family": "CRD_IV"},
        "FATF": {"name": "Financial Action Task Force", "jurisdiction": "GLOBAL", "family": "AML_CFT"},
        "SEC": {"name": "US Securities and Exchange Commission", "jurisdiction": "US", "family": "SEC"},
        "FED": {"name": "US Federal Reserve", "jurisdiction": "US", "family": "Basel_III"},
    }

    def extract_from_filename(self, filename: str) -> dict:
        """Parse structured metadata from filename."""
        stem = Path(filename).stem.upper()
        parts = stem.split("_")
        
        metadata = {
            "source": "UNKNOWN",
            "regulation_family": "UNKNOWN",
            "jurisdiction": "GLOBAL",
            "version": datetime.utcnow().strftime("%Y-%m-%d"),
            "doc_id": stem,
        }

        if parts and parts[0] in self.SOURCE_MAP:
            src = parts[0]
            info = self.SOURCE_MAP[src]
            metadata["source"] = src
            metadata["regulation_family"] = info["family"]
            metadata["jurisdiction"] = info["jurisdiction"]

        # Try to extract year
        for part in parts:
            if part.isdigit() and len(part) == 4 and 2000 <= int(part) <= 2099:
                metadata["version"] = f"{part}-01-01"
                break

        return metadata

    def extract_from_content(self, content: str, existing_meta: dict) -> dict:
        """Supplement metadata from document content (title, dates)."""
        # Extract title from first non-empty lines
        lines = [l.strip() for l in content.split('\n') if l.strip()]
        if lines:
            existing_meta["title"] = lines[0][:200]  # Cap at 200 chars

        # Look for "supersedes" / "replaces" references
        supersedes_match = re.search(
            r'(?:supersedes|replaces|amends)\s+(?:circular\s+)?([A-Z0-9/_\-\.]+)',
            content[:2000],
            re.IGNORECASE
        )
        if supersedes_match:
            existing_meta["supersedes"] = supersedes_match.group(1).upper()

        # Look for effective date
        date_match = re.search(
            r'effective\s+(?:from\s+)?(\d{1,2}[\s/\-]\w+[\s/\-]\d{4}|\w+\s+\d{1,2},?\s+\d{4})',
            content[:3000],
            re.IGNORECASE
        )
        if date_match:
            existing_meta["effective_date"] = date_match.group(1)

        return existing_meta


class DocumentLoader:
    """
    Main ingestion orchestrator.
    Parses documents, extracts metadata, and prepares them for chunking.
    """

    def __init__(self):
        self.parser = DocumentParser()
        self.meta_extractor = MetadataExtractor()
        self._processed_hashes: set = set()  # Deduplication

    def load_document(self, file_path: str, metadata_override: Optional[dict] = None) -> Optional[RegulatoryDocument]:
        """
        Load and parse a single regulatory document.
        Returns None if document is a duplicate (same hash already ingested).
        """
        file_path = str(file_path)
        filename = os.path.basename(file_path)
        ext = Path(file_path).suffix.lower()

        logger.info(f"Loading document: {filename}")

        # Parse content based on file type
        try:
            if ext == ".pdf":
                content = self.parser.parse_pdf(file_path)
            elif ext in (".docx", ".doc"):
                content = self.parser.parse_docx(file_path)
            elif ext in (".txt", ".md", ".html"):
                content = self.parser.parse_text(file_path)
            else:
                logger.warning(f"Unsupported file type: {ext} — skipping {filename}")
                return None
        except Exception as e:
            logger.error(f"Failed to parse {filename}: {e}")
            return None

        if not content or len(content.strip()) < 100:
            logger.warning(f"Document {filename} has insufficient content — skipping")
            return None

        # Deduplication check
        doc_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
        if doc_hash in self._processed_hashes:
            logger.info(f"Duplicate document detected: {filename} — skipping")
            return None
        self._processed_hashes.add(doc_hash)

        # Extract metadata
        metadata = self.meta_extractor.extract_from_filename(filename)
        metadata = self.meta_extractor.extract_from_content(content, metadata)
        if metadata_override:
            metadata.update(metadata_override)

        doc = RegulatoryDocument(
            content=content,
            doc_id=metadata["doc_id"],
            source=metadata["source"],
            regulation_family=metadata["regulation_family"],
            jurisdiction=metadata["jurisdiction"],
            version=metadata["version"],
            effective_date=metadata.get("effective_date"),
            supersedes=metadata.get("supersedes"),
            title=metadata.get("title", ""),
            file_path=file_path,
            doc_hash=doc_hash,
        )

        logger.info(
            f"Loaded: {doc.doc_id} | Source: {doc.source} | "
            f"Jurisdiction: {doc.jurisdiction} | Chars: {len(content)}"
        )
        return doc

    def load_directory(self, docs_dir: str, metadata_overrides: Optional[dict] = None) -> list[RegulatoryDocument]:
        """Load all supported documents from a directory."""
        docs_dir = Path(docs_dir)
        if not docs_dir.exists():
            raise FileNotFoundError(f"Documents directory not found: {docs_dir}")

        supported_extensions = {".pdf", ".docx", ".txt", ".md", ".html"}
        files = [f for f in docs_dir.iterdir() if f.suffix.lower() in supported_extensions]

        logger.info(f"Found {len(files)} documents to ingest from {docs_dir}")

        documents = []
        for file_path in sorted(files):  # Sort for deterministic processing order
            override = (metadata_overrides or {}).get(file_path.name, {})
            doc = self.load_document(str(file_path), override)
            if doc:
                documents.append(doc)

        logger.info(f"Successfully loaded {len(documents)}/{len(files)} documents")
        return documents

    
if __name__ == "__main__":
    import sys
    import argparse

    logging.basicConfig(level=logging.INFO)

    parser = argparse.ArgumentParser(description="Ingest regulatory documents")
    parser.add_argument("--docs-dir", default="sample_docs", help="Directory of regulatory documents")
    args = parser.parse_args()

    loader = DocumentLoader()
    docs = loader.load_directory(args.docs_dir)
    print(f"\n✅ Loaded {len(docs)} documents")
    for doc in docs:
        print(f"  - {doc.doc_id} ({doc.source}, {doc.jurisdiction}, {len(doc.content)} chars)")
